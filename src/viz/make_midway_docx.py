"""Generate docs/midway_summary.docx (partner-facing) — mirrors docs/midway_summary.md.
Run: python src/viz/make_midway_docx.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "midway_summary.docx"
ACCENT = RGBColor(0xB0, 0x2A, 0x2A); GREY = RGBColor(0x55, 0x55, 0x55)


def main():
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)

    def h(t, lvl=1):
        p = doc.add_heading(t, level=lvl)
        if lvl == 1:
            p.runs[0].font.color.rgb = ACCENT
        return p

    def bullets(items):
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    def table(headers, rows, bold0=False):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Light Grid Accent 1"
        for j, hd in enumerate(headers):
            r = t.rows[0].cells[j].paragraphs[0].add_run(hd); r.bold = True; r.font.size = Pt(9)
        for row in rows:
            c = t.add_row().cells
            for j, v in enumerate(row):
                rn = c[j].paragraphs[0].add_run(str(v)); rn.font.size = Pt(9)
                if bold0 and j == 0:
                    rn.bold = True

    title = doc.add_heading("CS2 Win-Probability Model — Midway Status", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run("For Leu · 2026-06-25 · de_inferno · arXiv + MLSA 2027"); r.font.color.rgb = GREY; r.font.size = Pt(9.5)

    p = doc.add_paragraph()
    p.add_run("TL;DR: ").bold = True
    p.add_run("4 feature pillars built; full 9-architecture matrix evaluated with bootstrap CIs + a "
              "thorough metric battery. BEST model = 4-model soft-vote ensemble 0.8531 AUC (best single: "
              "logistic on all-pillars EFB2 0.8515). Modeling phase essentially done and paper-ready; open "
              "work = firepower v2 (you), the 2026 holdout, and possibly more data.")

    h("1. Where the project stands")
    doc.add_paragraph("Predict P(CT win) every second of a round. 220 Tier-1 demos / 476,595 per-second "
                      "snapshots / 104 features; 5-fold GroupKFold by match. Honest headline: economy "
                      "dominates; map control + the defuse-race add small-but-significant lift; firepower is "
                      "the weakest pillar; deep models match (don't beat) a calibrated classical model at this "
                      "data size. Rigorous, calibrated, interpretable — a strong paper without a blockbuster AUC.")

    h("2. The four feature pillars")
    table(["Pillar", "Encodes", "Verdict"], [
        ["1 Economy/combat", "equip, HP, armor, alive, kits, score, time, bomb-planted",
         "Dominant; collapses to ~0.58 AUC in even rounds"],
        ["2 Map control (x3)", "Voronoi / grey (LOS+FOV+smoke) / territory (memory+decay)",
         "Voronoi real & significant (+0.003, +0.013 contested); grey/territory redundant"],
        ["3 Firepower (yours)", "HLTV Rating/ADR/KAST summed over alive + 1vN clutch",
         "Weakest pillar; has a count confound to fix in v2 (see s5)"],
        ["4 Tactical + bomb", "utility, AWP, entropy, zones, bomb geometry, defuse-race",
         "Defuse-race = best single addition since economy"],
    ], bold0=True)

    h("3. The model matrix (5-fold OOF, B=500 bootstrap CIs)")
    table(["model", "AUC", "95% CI", "ECE", "cAUC"], [
        ["logreg EFB2 (best single)", "0.8515", "(0.844,0.858)", "0.016", "0.596"],
        ["xgb EFB2", "0.8498", "(0.843,0.857)", "0.013", "0.591"],
        ["lgbm EFB2", "0.8498", "-", "0.012", "0.593"],
        ["catboost EFB2", "0.8483", "-", "0.018", "0.586"],
        ["RF EFB2 (cautionary)", "0.8442", "-", "-", "0.574"],
        ["TCN (Betty GPU)", "0.8489", "(0.842,0.856)", "0.012", "0.572"],
        ["Transformer (Betty GPU)", "0.8473", "(0.840,0.854)", "0.009", "0.568"],
        ["GAT raw trajectories (GPU)", "0.8465", "(0.840,0.853)", "0.014", "0.576"],
        ["SOFT-VOTE (4) — BEST", "0.8531", "(0.846,0.860)", "0.009", "0.589"],
        ["logistic-stack", "0.8529", "-", "0.042", "0.590"],
    ], bold0=True)
    doc.add_paragraph("Every point estimate is inside the others' 95% CIs -> statistical dead heat. No deep "
                      "model beats classical; the ensemble wins by combining diverse learners.")

    h("4. Key scientific findings")
    bullets([
        "Contested-AUC (novel lens): economy collapses to ~0.58 in even rounds; spatial pillars earn value there.",
        "Map control (Voronoi) robustly significant across all models; grey/territory redundant (clean ablation).",
        "Defuse-race (fuse-left - run-time - defuse-time) = #8 feature overall; post-plant log-loss -7-8%.",
        "Deep models tie, don't beat: momentum (TCN) and raw trajectories (GAT) don't help at 220 matches.",
        "The accuracy gain is RESOLUTION not calibration (Brier decomposition): features add skill, calibration stays ~perfect.",
        "Honest probabilities + comebacks: ECE<0.02; says 10% -> side wins 6.8%; winner written off (<=10%) in 7.2% of rounds, calibrated.",
    ])

    h("5. Firepower (your pillar) — v1 -> v2 (both benchmarked)")
    doc.add_paragraph("v1 (sum-based, 9 feats): summed HLTV Rating/ADR/KAST + 1vN clutch, joined (steamid, year). "
                      "F-A significant only on logreg; EF-E ~= 0; value conditional (contested +0.007). KEY ISSUE = a "
                      "count confound: firepower_rating_diff was perm-importance #1 but 0.988-correlated with the "
                      "player-count advantage (Rating is SUMMED -> ~99% a count proxy, not skill).")
    doc.add_paragraph("v2 (your redesign, commit fc4f719, 20 feats): side-aware CT/T stats (player_stats_sided.csv), "
                      "per-player gates (lone survivor->Clutch; teammates->Entry/Trading; Opening only at 5v5), AWP-holder "
                      "Sniping flag, grenade-$-weighted Utility. Benchmarked on the same 220 demos / 5-fold OOF / B=500.")
    table(["metric", "v1", "v2"], [
        ["logreg F-A", "+0.0018 (sig)", "+0.0022 (sig)"],
        ["logreg contested-AUC (F)", "0.593", "0.603 (up)"],
        ["logreg EFB2 (best overall)", "0.8515", "0.8519 (up, study best)"],
        ["xgb/lgbm/catboost EF-E", "~0 (ns)", "negative (catboost -0.0026 sig)"],
        ["count confound (rating-diff vs count)", "0.988", "0.987 (persists)"],
    ], bold0=True)
    bullets([
        "v2 HELPED the linear/headline model: best contested-AUC (0.603) and best overall AUC (logreg EFB2 0.8519) in the study.",
        "v2 HURT the tree models (EF < E; catboost significantly): the 20 sparse NaN-gated features overfit the GBMs.",
        "The count confound is STILL there (v2 kept sums). Side-awareness/gating fixed a different axis, not this one.",
    ])
    p = doc.add_paragraph(); p.add_run("Suggested v3 (open item): ").bold = True
    bullets([
        "Use AVERAGE rating per alive player (per-capita), not the sum -> finally decouple skill from count.",
        "Prune the sparse gated features that hurt GBMs; keep what permutation importance likes (rating->avg, adr, t_trading, AWP sniping).",
        "Keep v2's good parts: side-awareness + year-aware (steamid, year) join.",
    ])

    h("5b. 2026 OUT-OF-TIME HOLDOUT — the big new result (+ action item for you)")
    doc.add_paragraph("Trained on all 220 demos (2024-25), evaluated ONCE on 27 fresh 2026 Inferno matches "
                      "(55,271 snapshots; base rate shifts 0.445 -> 0.512).")
    table(["", "in-time (OOF)", "out-of-time 2026", "delta"], [
        ["lgbm EB2 (no firepower)", "0.8493", "0.8501", "+0.0008"],
        ["xgb EB2", "0.8489", "0.8497", "+0.0007"],
        ["xgb E", "0.8476", "0.8476", "+0.0000"],
        ["logreg EFB2 (WITH firepower)", "0.8519 (best in-sample)", "0.8236", "-0.0283"],
        ["rf EFB2", "0.8446", "0.8159", "-0.0287"],
    ], bold0=True)
    bullets([
        "THE CORE MODEL GENERALISES: without firepower, out-of-time ~= in-time (some better). Contested-AUC even "
        "IMPROVES (0.590 -> 0.64-0.65). Economy + map control + tactical + defuse-race TRANSFER.",
        "FIREPOWER DOES NOT TRANSFER: EFB2 collapses; ECE 0.016 -> 0.071; calibration intercept -0.36. "
        "The best in-sample model became the WORST out-of-time.",
        "DIAGNOSIS (data-coverage gap, fixable): 0 of 27 2026 demos are in demo_year_map.csv (-> 2024 stats used), "
        "and ~30% of 2026 players have no HLTV stats (5v5 mean ct_rating_sum: 5.28 train vs 3.66 on 2026; "
        "11.8% of 2026 5v5 snapshots < 3 vs 0.0% in training). Model reads corrupted firepower as 'weak team'.",
    ])
    p = doc.add_paragraph(); p.add_run("ACTION FOR YOU (Leu): ").bold = True
    bullets([
        "Add the 27 2026 demos to configs/demo_year_map.csv with year=2026.",
        "Scrape 2026 HLTV stats for the missing 2026 players into player_stats_sided.csv.",
        "Then we re-run the holdout as a SEPARATE, DISCLOSED evaluation (the first run stands as the honest touch-once result).",
    ])
    doc.add_paragraph("Lesson (a real paper contribution): a skill-prior pillar creates an INFERENCE-TIME DATA "
                      "DEPENDENCY the other pillars don't have (they're computed from the demo itself). On "
                      "out-of-time evidence the recommended production model is EB2 (no firepower).")

    h("6. Evaluation framework")
    doc.add_paragraph("Primary: AUC, log-loss, Brier. Complementary: ECE, BSS, contested-AUC. Extended (new): "
                      "Brier decomposition (reliability/resolution/uncertainty), sharpness, bin-free calibration "
                      "(slope+intercept, adaptive-ECE, KS-cal), and a comeback/tail honesty diagnostic. Uncertainty: "
                      "match-level block bootstrap (B=500) CIs on every metric; multi-seed std for deep models. "
                      "Interpretation: coefficients, permutation importance, SHAP. Full glossary in docs/metrics.md. "
                      "Standing rule: every new model/method runs the same interp+uncertainty+calibration battery.")

    h("7. Infrastructure")
    table(["Local (laptop)", "Betty (GPU)"], [
        ["5 classical models, ensemble, ALL metrics, feature assembly", "TCN, Transformer, GAT (deep models)"],
    ])

    h("8. Next steps")
    bullets([
        "Firepower v2 (Leu): average-rating-per-player fix (s5) — highest near-term ROI for this pillar.",
        "2026 out-of-time holdout (Henry): parse ~15-20 fresh 2026 Inferno demos, run frozen EFB2 once — last validity gate.",
        "More data (together): ~2-5x matches and/or multi-map -- the real lever for spatial/deep models to surpass classical.",
        "Paper draft: methods + results complete (docs/results_checkpoint.md + docs/metrics.md).",
    ])

    h("9. Repo map")
    table(["area", "files"], [
        ["Features", "src/features/{economy,mapcontrol,positional,bomb,firepower,assemble,build_trajectory_dataset}.py"],
        ["Classical/eval", "src/models/{train_pipeline,calibration,conditional_analysis,logistic_coefficients,"
                           "permutation_importance,shap_analysis,extended_metrics,ensemble_oof}.py"],
        ["Deep (Betty)", "src/models/deep/{tcn,gat,transformer}.py + jobs/*.sh + docs/cluster_runbook.md"],
        ["Docs", "docs/{results_checkpoint,metrics,methodology,firepower_pillar,midway_summary}.md"],
    ], bold0=True)

    OUT.parent.mkdir(parents=True, exist_ok=True); doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
