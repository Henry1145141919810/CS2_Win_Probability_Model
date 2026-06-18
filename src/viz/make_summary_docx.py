"""Generate the partner-facing Word summary (docs/two_day_summary.docx).

Mirrors docs/two_day_summary.md but as a styled .docx (headings, tables, bullets) for easy
sharing. Run: python src/viz/make_summary_docx.py
"""
from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "two_day_summary.docx"
ACCENT = RGBColor(0xB0, 0x2A, 0x2A)
GREY = RGBColor(0x55, 0x55, 0x55)


def main():
    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)

    def h(txt, level=1):
        p = doc.add_heading(txt, level=level)
        if level == 1:
            p.runs[0].font.color.rgb = ACCENT
        return p

    def bullets(items):
        for it in items:
            doc.add_paragraph(it, style="List Bullet")

    def table(headers, rows, bold_first_col=False):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        for j, hd in enumerate(headers):
            c = t.rows[0].cells[j].paragraphs[0]
            r = c.add_run(hd); r.bold = True; r.font.size = Pt(9.5)
        for row in rows:
            cells = t.add_row().cells
            for j, val in enumerate(row):
                par = cells[j].paragraphs[0]
                rn = par.add_run(str(val)); rn.font.size = Pt(9.5)
                if bold_first_col and j == 0:
                    rn.bold = True
        return t

    # ---- Title ----
    title = doc.add_heading("CS2 Win-Probability Model", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Two-Day Progress Summary"); r.bold = True; r.font.size = Pt(14)
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("For the project partner  ·  Henry (UPenn WSABI)  ·  2026-06-16 → 2026-06-18\n"
                      "de_inferno  ·  live per-second round win-probability  ·  arXiv + MLSA 2027")
    rm.font.size = Pt(9.5); rm.font.color.rgb = GREY
    doc.add_paragraph()

    h("1. One-paragraph recap")
    doc.add_paragraph(
        "We are building a live, per-second CS2 round win-probability model for de_inferno that "
        "augments the established economy baseline (Xenopoulos/ESTA-style) with novel spatial "
        "feature pillars — Voronoi map control, line-of-sight/territory control, and tactical "
        "readiness — evaluated rigorously across model architectures. Over these two days the "
        "project went from a small pilot to an at-scale, fully-evaluated study: 220 pro demos / "
        "476,595 per-second snapshots, three competing map-control formulations, a five-"
        "architecture model matrix with bootstrap confidence intervals, calibration testing, "
        "interpretability, and qualitative case studies. The headline scientific result is "
        "honest and defensible: map control adds a small but statistically robust lift overall "
        "(~+0.003 AUC), concentrated in genuinely contested rounds (~+0.013) where the economy "
        "baseline collapses to a coin flip.")

    h("2. Dataset & pipeline")
    bullets([
        "220 Tier-1 de_inferno demos, de-duplicated and tier-filtered. Parsed with awpy v2.0.2, "
        "tickrate forced to 64.",
        "476,595 snapshots, one row per second from freeze-end to round end; label ct_won "
        "(base rate 0.445). Grouped by match_id for cross-validation (never split a match).",
        "Five parquet channels per demo (ticks/kills/rounds/bomb/smokes). Memory-guarded "
        "parsing (~6.7 GB/demo peak on a 16.8 GB laptop).",
        "Pillars implemented: (1) economy/combat, (2) map control — 3 variants, (4) tactical + "
        "bomb/rotation. Pillar (3) firepower is the main remaining gap.",
    ])

    h("3. The three map-control models (core contribution)")
    table(["Model", "Definition", "Verdict"], [
        ["Voronoi", "Each nav area → nearest living player's team (area-weighted). Greedy/total.",
         "Best single predictor; KEPT. Robustly significant in linear and nonlinear models."],
        ["Grey (LOS+FOV+smoke)", "Area held only if a living player is in range AND has line-of-sight "
         "AND is facing it AND it isn't smoked. ~60-80% of map is 'grey' at any instant.",
         "NEGATIVE result. Physically faithful but flickers tick-to-tick; does not beat Voronoi."],
        ["Territory (memory+decay)", "Grey model WITH memory: cleared space stays a team's for 15 s "
         "without re-checking; FOV only gates re-acquiring. Stateful per round.",
         "Recovers Voronoi-level predictiveness."],
    ], bold_first_col=True)
    p = doc.add_paragraph()
    p.add_run("Key insight — 'realistic ≠ predictive unless stabilized.' ").bold = True
    p.add_run("The instantaneous, physically correct sightline model is a worse predictor than "
              "crude proximity, because round outcomes depend on stable territorial state, not on "
              "where 10 crosshairs point this tick. Memory/decay fixes the flicker and restores the "
              "signal — a clean, publishable finding. All three models are retained for the ablation.")

    h("4. Model matrix — 5 architectures × {A, E, ET}")
    doc.add_paragraph("5-fold GroupKFold OOF; primary metric AUC; B=500 match-level block bootstrap. "
                      "A = economy only; E = economy + Voronoi + tactical/bomb; ET = E + territory.")
    table(["Model", "A", "E", "ET", "Lift E−A (95% CI)", "ECE (E)"], [
        ["Logistic", "0.8465", "0.8493", "0.8493", "+0.0028 (+0.0014, +0.0042)", "0.016"],
        ["XGBoost (tuned)", "0.8443", "0.8476", "0.8480", "+0.0034 (+0.0020, +0.0047)", "0.012"],
        ["LightGBM", "0.8448", "0.8479", "0.8476", "+0.0031 (+0.0018, +0.0042)", "0.014"],
        ["CatBoost", "0.8448", "0.8478", "0.8474", "+0.0030 (+0.0018, +0.0043)", "0.011"],
        ["Random Forest", "0.8228", "0.8426", "0.8428", "+0.0198 (+0.0165, +0.0228)", "0.011"],
    ], bold_first_col=True)
    bullets([
        "Every model, every feature set: the control lift is statistically significant (all "
        "bootstrap CIs exclude 0; DeLong p ≈ 0).",
        "The four well-specified models agree on the honest magnitude: ~+0.003 AUC. ET ≈ E "
        "everywhere → territory is redundant with Voronoi.",
        "Logistic regression is the best model (0.8493) and ties the GBMs → the signal is "
        "essentially linear; the gradient-boosted models are interchangeable.",
        "Random Forest is the cautionary cell: its economy baseline underfits (0.8228, ECE 0.042), "
        "so control 'helps' a huge +0.0198 and fixes calibration — a weak-baseline artifact, not "
        "extra signal. This is why we report the lift across architectures, not from one model.",
    ])

    h("5. Metrics — primary + a metric that tells the honest story")
    bullets([
        "Primary (literature-comparable): AUC-ROC, log-loss, Brier.",
        "Complementary (now standard harness outputs): ECE (calibration), BSS (Brier Skill "
        "Score), and contested-AUC (cAUC) — AUC on genuinely contested snapshots (equal players "
        "alive AND |equip diff| ≤ 1500; 58,368 snapshots).",
    ])
    p = doc.add_paragraph()
    p.add_run("Why cAUC matters: ").bold = True
    p.add_run("overall AUC is dominated by easy lopsided snapshots that economy already nails, "
              "diluting the spatial signal and making the aggregate lift look tiny (+0.003). A "
              "better metric will NOT inflate that — it is genuinely small on average. But cAUC "
              "reports the lift where it matters: in even rounds the economy baseline's own AUC "
              "collapses from ~0.85 to ~0.58 (near coin-flip). That is the honest and stronger "
              "framing for the paper.")

    h("6. Where map control matters — conditional analysis")
    bullets([
        "Economy baseline collapses in even rounds: all 0.832 → equal-alive 0.686 → even-econ "
        "0.674 → equal-alive AND even-econ 0.578 (~coin flip).",
        "Map-control lift is larger where it matters: equal-alive E−A = +0.0131 vs +0.009 overall; "
        "even-econ +0.0107; pre-plant +0.0103.",
        "The signal is nonlinear in the hardest subsets — XGBoost extracts the contested lift, "
        "logistic gets ~0 there (logistic wins easy snapshots, XGBoost wins where it's hard).",
    ])

    h("7. Interpretability — what the model learned")
    doc.add_paragraph("Standardized logistic coefficients (z-scored inputs ⇒ magnitude = effect "
                      "size, sign = direction toward CT win). Signs all sensible; the relative "
                      "scale is the real story — economy dominates, map control is real but "
                      "second-order:")
    bullets([
        "Economy/combat dominate (set A, ~0.5–0.9): t_players_alive −0.89, ct_health_total +0.74, "
        "ct_equipment_value +0.72, t_equipment_value −0.61, bomb_planted −0.51; plus the single "
        "largest term min_ct_dist_to_bomb −0.78 (CTs far from bomb = retake = bad).",
        "Map control fit on its own ≈ +0.11 (Voronoi control_deficit) — the strongest "
        "non-economy/non-bomb signal, but ~5–8× smaller than economy.",
        "Voronoi vs territory, fit separately: Voronoi deficit +0.115 vs territory deficit +0.066 "
        "— territory is directionally the same but about half as strong, and the two are only "
        "moderately correlated (r = 0.47). Territory is a related-but-noisier proxy, redundant "
        "with Voronoi at the AUC level (ET ≈ E), not an identical or stronger signal.",
        "Caveat: coefficients inside the full model are collinearity-inflated (control_deficit "
        "shows +0.27 there only because a duplicate column splits the weight) — the standalone "
        "fits above are the trustworthy effect sizes.",
    ])

    h("8. Qualitative case studies — map control flipping the call")
    p = doc.add_paragraph()
    p.add_run("Headline example — b8 vs FlyQuest (map 3), round 12, a 4-v-4 with even economy: ").bold = True
    p.add_run("the economy-only model hugs 0.50 (a coin flip) for the entire round, but the "
              "map-control model consistently and correctly pulls toward T, who won (peak shift "
              "−22% at t+50 s). This single figure makes the contested-AUC number concrete: economy "
              "is blind in even rounds; control sees the read. Two more examples are included, each "
              "with a win-probability timeline plus the 3-model map at the key moment.")

    h("9. Calibration — are the probabilities honest?")
    doc.add_paragraph("A reliability diagram with per-bin bootstrap error bars and a bootstrap 95% "
                      "CI on ECE (resampling whole matches, B=200). Result: all non-RF models are "
                      "well-calibrated (ECE < 0.02) — 'the model said 70%' really means CT won ~70% "
                      "of the time. This is now a standard, repeatable check, not a one-off.")

    h("10. Evaluation protocol")
    bullets([
        "5-fold GroupKFold by match — never split a match (avoids leakage that inflates AUC).",
        "DeLong's test for each feature set vs the economy baseline on identical OOF predictions.",
        "Match-level block bootstrap (matches are the independent unit), B=500; a difference CI "
        "excluding 0 ⇒ significant.",
        "Time-window analysis at 5/10/15/20/25 s for the control-signal-emergence story.",
        "Held out for the end: a 2026 out-of-time test set, touched once, kept uncontaminated.",
    ])

    h("11. Repo map & reproduction")
    table(["Area", "Files"], [
        ["Features", "src/features/{economy,mapcontrol,positional,bomb,assemble}.py, visibility.py"],
        ["Models / eval", "src/models/{train_pipeline,calibration,conditional_analysis,tune_xgb,"
                          "logistic_coefficients}.py"],
        ["Visualization", "src/viz/{mapcontrol_viz,control_shift_examples,winprob_chart}.py"],
        ["Docs", "docs/methodology.md (full protocol + every result), docs/map_control_models.md"],
    ], bold_first_col=True)
    pr = doc.add_paragraph()
    pr.add_run("Reproduce the headline table:  ").bold = True
    code = pr.add_run("python src/models/train_pipeline.py --models logreg,xgb,rf,lgbm,catboost "
                      "--sets A,E,ET --bootstrap 500")
    code.font.name = "Consolas"; code.font.size = Pt(9)
    note = doc.add_paragraph()
    note.add_run("Note: figures live in outputs/figures/ and are git-ignored (regenerable from the "
                 "scripts above). The parsed dataset and demos are shared separately, not in git. "
                 "Ask if you want a specific figure committed.").italic = True

    h("12. Open items / next steps")
    bullets([
        "Pillar 3 — firepower (per-player rating/aim stats): clearest remaining lever; needs an "
        "HLTV per-player scrape + name mapping (manual, ToS-limited).",
        "Deep models on E — GAT / TCN / Transformer / ensemble (cloud A100), mean ± std over 20 "
        "seeds; the laptop covers the classical matrix.",
        "2026 out-of-time holdout — final generalization check.",
        "Honest positioning: small aggregate lift, robustly significant, concentrated in contested "
        "rounds; the three-control-model ablation + calibration + contested-AUC story are the "
        "methodological contributions even if the raw AUC gain is modest.",
    ])

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
